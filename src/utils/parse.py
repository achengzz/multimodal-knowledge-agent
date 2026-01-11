from pathlib import Path

import pandas as pd
from PyPDF2 import  PdfReader
from bs4 import BeautifulSoup
from docx import Document

from src.utils.chunk import chunk_by_zh_chars


def read_text(path):
    with open(path, 'r', encoding='utf-8') as f:
        return f.read()

def read_pdf(path):
    text = ""
    with open(path, 'rb') as file:
        pdf_reader = PdfReader(file)
        num_pages = len(pdf_reader.pages)
        for page_num in range(num_pages):
            page = pdf_reader.pages[page_num]
            text += page.extract_text().replace('\n', '')
    return text

# def read_pdf(path):
#
#     reader = PdfReader(path)
#     text=[]
#     for i, page in enumerate(reader.pages):
#         page_text=page.extract_text()
#         if page_text:
#             text.append(f"-- 页码_{i+1} --\n{page_text}")
#
#     return "\n\n".join(text)


def read_docx(path):
    doc=Document(path)
    paragraphs = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            paragraphs.append(paragraph.text)

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)

    return "\n\n".join(paragraphs)

def read_csv(path):
    df=pd.read_csv(path)
    text=[f"共{len(df)}行 {len(df.columns)}列"]
    text.append(f"列名：{', '.join(df.columns.tolist())}")
    text.append("\n数据：")
    text.append(df.to_markdown(index=False))

    return "\n".join(text)


def read_xlsx(path):
    xlsx=pd.ExcelFile(path)
    text=[]
    for sheet in xlsx.sheet_names:
        df=pd.read_excel(xlsx,sheet_name=sheet)
        text.append(f"\n### sheet {sheet}")
        text.append(f"rows: {len(df)}, columns: {len(df.columns)}")
        text.append(df.to_markdown(index=False))

    return "\n".join(text)

def read_html(path):
    with open(path, 'r',encoding='utf-8') as f:
        content=f.read()
    soup=BeautifulSoup(content, 'html.parser')

    for element in soup(["script","style","nav","footer","header"]):
        element.decompose()

    text=soup.get_text(separator="\n")

    lines=[line.strip() for line in text.splitlines() if line.strip()]
    return "\n".join(lines)


def parse_file(path):

    path = Path(path)

    filename = path.name.lower()

    if filename.endswith('.pdf'):
        return read_pdf(path)
    elif filename.endswith('.docx'):
        return read_docx(path)
    elif filename.endswith('.xlsx'):
        return read_xlsx(path)
    elif filename.endswith('.csv'):
        return read_csv(path)
    elif filename.endswith('.html'):
        return read_html(path)
    else:
        return read_text(path)


